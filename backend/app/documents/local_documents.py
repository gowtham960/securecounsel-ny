import csv
from datetime import datetime, timezone
from io import BytesIO, StringIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from app.ingestion.local_index import chunk_text
from app.retrieval.dense_qdrant import upsert_chunks


UPLOAD_ROOT = Path("data") / "uploaded_docs"
ORIGINAL_UPLOAD_ROOT = Path("data") / "uploaded_originals"

DOCUMENT_REGISTRY: list[dict] = []

SUPPORTED_EXTENSIONS = {
    ".txt",
    ".pdf",
    ".docx",
    ".xlsx",
    ".csv",
}


def utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat().replace(
        "+00:00", "Z"
    )


def _rebuild_registry_from_disk():
    """Scan uploaded_docs on disk and rebuild DOCUMENT_REGISTRY at startup."""
    if not UPLOAD_ROOT.exists():
        return

    for matter_folder in UPLOAD_ROOT.iterdir():
        if not matter_folder.is_dir():
            continue

        matter_id = matter_folder.name

        for txt_file in matter_folder.glob("*.txt"):
            try:
                text = txt_file.read_text(encoding="utf-8")
            except Exception:
                continue

            def _field(name: str, t: str = text) -> str:
                for line in t.splitlines():
                    if line.startswith(f"{name}:"):
                        return line.replace(f"{name}:", "", 1).strip()
                return ""

            doc_id = _field("Document ID")
            if not doc_id:
                continue

            if any(d["document_id"] == doc_id for d in DOCUMENT_REGISTRY):
                continue

            DOCUMENT_REGISTRY.append({
                "document_id": doc_id,
                "matter_id": matter_id,
                "filename": _field("Title"),
                "stored_path": str(txt_file),
                "uploaded_by": _field("Uploaded By"),
                "uploaded_at": _field("Uploaded At"),
                "content_type": "",
                "status": "restored_from_disk",
                "source_extension": _field("Original File Type"),
            })


# Rebuild registry from disk at import time (fixes restart bug)
_rebuild_registry_from_disk()


def ensure_upload_folder(matter_id: str) -> Path:
    folder = UPLOAD_ROOT / matter_id
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def ensure_original_upload_folder(matter_id: str) -> Path:
    folder = ORIGINAL_UPLOAD_ROOT / matter_id
    folder.mkdir(parents=True, exist_ok=True)
    return folder


def sanitize_filename(filename: str) -> str:
    safe = filename.replace("\\", "_").replace("/", "_").strip()

    if not safe:
        safe = "uploaded_document.txt"

    return safe


def _extract_text_from_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError("Uploaded .txt or .csv file must be UTF-8 encoded.") from exc


def _extract_text_from_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(content))
    except Exception as exc:
        raise ValueError("Could not read the uploaded PDF file.") from exc

    parts = []

    for page_index, page in enumerate(reader.pages, start=1):
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""

        page_text = page_text.strip()

        if page_text:
            parts.append(f"Page {page_index}\n{page_text}")

    text = "\n\n".join(parts).strip()

    if not text:
        raise ValueError(
            "No searchable text could be extracted from this PDF. "
            "Scanned PDFs will need OCR support later."
        )

    return text


def _extract_text_from_docx(content: bytes) -> str:
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise ValueError("Could not read the uploaded DOCX file.") from exc

    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        parts.append(f"Table {table_index}")

        for row_index, row in enumerate(table.rows, start=1):
            row_values = []

            for cell in row.cells:
                value = cell.text.strip()
                row_values.append(value)

            if any(row_values):
                parts.append(f"Row {row_index}: " + " | ".join(row_values))

    text = "\n".join(parts).strip()

    if not text:
        raise ValueError("No searchable text could be extracted from this DOCX file.")

    return text


def _extract_text_from_xlsx(content: bytes) -> str:
    try:
        workbook = load_workbook(BytesIO(content), data_only=True, read_only=True)
    except Exception as exc:
        raise ValueError("Could not read the uploaded XLSX file.") from exc

    parts = []

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        parts.append(f"Sheet: {sheet_name}")

        for row_index, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            values = []

            for value in row:
                if value is None:
                    values.append("")
                else:
                    values.append(str(value).strip())

            if any(values):
                parts.append(f"Row {row_index}: " + " | ".join(values))

    text = "\n".join(parts).strip()

    if not text:
        raise ValueError("No searchable text could be extracted from this XLSX file.")

    return text


def _extract_text_from_csv(content: bytes) -> str:
    raw_text = _extract_text_from_txt(content)

    try:
        reader = csv.reader(StringIO(raw_text))
    except Exception as exc:
        raise ValueError("Could not read the uploaded CSV file.") from exc

    parts = []

    for row_index, row in enumerate(reader, start=1):
        values = [value.strip() for value in row]

        if any(values):
            parts.append(f"Row {row_index}: " + " | ".join(values))

    text = "\n".join(parts).strip()

    if not text:
        raise ValueError("No searchable text could be extracted from this CSV file.")

    return text


def extract_text_from_upload(filename: str, content: bytes) -> tuple[str, str]:
    safe_name = sanitize_filename(filename)
    extension = Path(safe_name).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Unsupported file type. Supported types: .txt, .pdf, .docx, .xlsx, .csv."
        )

    if extension == ".txt":
        return _extract_text_from_txt(content), "text/plain"

    if extension == ".pdf":
        return _extract_text_from_pdf(content), "application/pdf"

    if extension == ".docx":
        return _extract_text_from_docx(content), (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    if extension == ".xlsx":
        return _extract_text_from_xlsx(content), (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    if extension == ".csv":
        return _extract_text_from_csv(content), "text/csv"

    raise ValueError(
        "Unsupported file type. Supported types: .txt, .pdf, .docx, .xlsx, .csv."
    )


def save_uploaded_text_document(
    matter_id: str,
    uploaded_by: str,
    filename: str,
    content: bytes,
) -> dict:
    safe_name = sanitize_filename(filename)
    extension = Path(safe_name).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            "Unsupported file type. Supported types: .txt, .pdf, .docx, .xlsx, .csv."
        )

    extracted_text, content_type = extract_text_from_upload(
        filename=safe_name,
        content=content,
    )

    document_id = "doc_upload_" + uuid4().hex[:12]

    original_folder = ensure_original_upload_folder(matter_id)
    extracted_folder = ensure_upload_folder(matter_id)

    original_filename = f"{document_id}_{safe_name}"
    original_path = original_folder / original_filename
    original_path.write_bytes(content)

    extracted_filename = f"{document_id}_{Path(safe_name).stem}.txt"
    extracted_path = extracted_folder / extracted_filename

    searchable_text = (
        f"Document ID: {document_id}\n"
        f"Matter ID: {matter_id}\n"
        f"Firm ID: firm_demo\n"
        f"Title: {safe_name}\n"
        f"Original File Type: {extension}\n"
        f"Uploaded By: {uploaded_by}\n"
        f"Uploaded At: {utc_now_iso()}\n\n"
        f"{extracted_text}"
    )

    extracted_path.write_text(searchable_text, encoding="utf-8")

    # Index into Qdrant for semantic search
    try:
        text_chunks = chunk_text(extracted_text)
        qdrant_chunks = [
            {
                "text": chunk,
                "metadata": {
                    "document_id": document_id,
                    "chunk_id": f"{document_id}_chunk_{i+1:03d}",
                    "firm_id": "firm_demo",
                    "matter_id": matter_id,
                    "collection": "uploaded_matter_docs",
                    "source_type": "uploaded_file",
                    "title": safe_name,
                    "citation": f"Uploaded document: {safe_name}",
                }
            }
            for i, chunk in enumerate(text_chunks)
        ]
        upsert_chunks(qdrant_chunks)
    except Exception as e:
        print(f"[qdrant] Failed to index {safe_name}: {e}")

    document = {
        "document_id": document_id,
        "matter_id": matter_id,
        "filename": safe_name,
        "original_path": str(original_path),
        "stored_path": str(extracted_path),
        "uploaded_by": uploaded_by,
        "uploaded_at": utc_now_iso(),
        "content_type": content_type,
        "status": "uploaded_and_extracted",
        "size_bytes": len(content),
        "extracted_text_bytes": len(searchable_text.encode("utf-8")),
        "source_extension": extension,
    }

    DOCUMENT_REGISTRY.append(document)

    return document


def list_documents_for_matter(matter_id: str) -> list[dict]:
    return [
        document
        for document in DOCUMENT_REGISTRY
        if document["matter_id"] == matter_id
    ]